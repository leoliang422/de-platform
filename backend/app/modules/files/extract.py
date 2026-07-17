"""投稿文件解析：把上传的本地文件转成可插入投稿正文的 Markdown 片段。

分层：
- 文本类（txt/md/csv/json）直接解码，无需外部依赖（真实可用）。
- 图片直接以 Markdown 图片语法内联，前端即可展示（真实可用）。
- Word(.docx) / PDF 用本地解析库（python-docx / pdfplumber）抽取纯文本，**不依赖大模型**：
    - ``LocalDocExtractor``（默认）：真实抽取文字；抽不出（旧版 .doc、扫描件、解析失败）时
      回退占位提示 + 下载链接。
    - ``LLMExtractor``：在本地抽取基础上再交大模型归一为规范 Markdown 的骨架（占位）；
      未接入时工厂回退本地解析，因此不影响现有功能。
"""

from __future__ import annotations

import logging
from dataclasses import dataclass
from io import BytesIO
from typing import Protocol, runtime_checkable

from app.core.config import get_settings

logger = logging.getLogger(__name__)

# content_type -> 扩展名。分三类，决定不同的处理路径。
TEXT_TYPES: dict[str, str] = {
    "text/plain": ".txt",
    "text/markdown": ".md",
    "text/x-markdown": ".md",
    "text/csv": ".csv",
    "application/json": ".json",
}
IMAGE_TYPES: dict[str, str] = {
    "image/png": ".png",
    "image/jpeg": ".jpg",
    "image/gif": ".gif",
    "image/webp": ".webp",
}
DOCX_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PDF_TYPE = "application/pdf"
DOCUMENT_TYPES: dict[str, str] = {
    DOCX_TYPE: ".docx",
    "application/msword": ".doc",  # 旧版二进制 Word，本地库无法解析 → 占位
    PDF_TYPE: ".pdf",
}
ALLOWED_TYPES: dict[str, str] = {**TEXT_TYPES, **IMAGE_TYPES, **DOCUMENT_TYPES}


@dataclass
class ExtractOutput:
    text: str  # 可直接插入投稿正文的 Markdown
    placeholder: bool  # True 表示未真正解析出内容（仅占位提示）


def _placeholder(filename: str, url: str | None, note: str) -> ExtractOutput:
    link = f"[{filename}]({url})" if url else filename
    return ExtractOutput(text=f"> 已上传文件：{link}\n>\n> {note}\n", placeholder=True)


def _parse_docx(data: bytes) -> str:
    from docx import Document  # 延迟导入，避免无谓加载

    doc = Document(BytesIO(data))
    lines: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # Word 标题样式 → Markdown 标题，尽量保留层级
        style = (para.style.name or "").lower() if para.style else ""
        if style.startswith("heading"):
            level = "".join(ch for ch in style if ch.isdigit()) or "2"
            lines.append(f"{'#' * min(int(level), 6)} {text}")
        else:
            lines.append(text)
    # 表格按管道表拼接
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            if any(cells):
                lines.append("| " + " | ".join(cells) + " |")
    return "\n\n".join(lines).strip()


def _parse_pdf(data: bytes) -> str:
    import pdfplumber

    parts: list[str] = []
    with pdfplumber.open(BytesIO(data)) as pdf:
        for page in pdf.pages:
            text = (page.extract_text() or "").strip()
            if text:
                parts.append(text)
    return "\n\n".join(parts).strip()


@runtime_checkable
class DocumentExtractor(Protocol):
    name: str

    async def extract(
        self, *, data: bytes, filename: str, content_type: str, url: str | None
    ) -> ExtractOutput: ...


class LocalDocExtractor:
    """用本地解析库抽取 Word/PDF 文字，不依赖大模型。"""

    name = "local"

    async def extract(
        self, *, data: bytes, filename: str, content_type: str, url: str | None
    ) -> ExtractOutput:
        try:
            if content_type == DOCX_TYPE:
                text = _parse_docx(data)
            elif content_type == PDF_TYPE:
                text = _parse_pdf(data)
            else:
                # 旧版 .doc 等本地库不支持
                return _placeholder(
                    filename, url, "暂不支持旧版 .doc 解析，请另存为 .docx 或手动粘贴正文。"
                )
        except Exception as exc:  # noqa: BLE001 - 解析失败不应阻断投稿
            logger.warning("文件 %s 解析失败：%s", filename, exc, exc_info=True)
            return _placeholder(filename, url, "文件解析失败，请检查文件或手动粘贴正文。")

        if not text:
            return _placeholder(
                filename,
                url,
                "未从文件中提取到文字（可能是扫描件/纯图片 PDF），请手动补充或改用图片上传。",
            )
        return ExtractOutput(text=text, placeholder=False)


class LLMExtractor:
    """本地抽取 + 大模型归一为规范 Markdown 的骨架（占位）。

    真实实现：先用 ``LocalDocExtractor`` 抽文字，再送 LLM 归一；图片可走多模态 OCR。
    未接入（``is_configured`` 为 False）时工厂回退本地解析，不影响现有功能。
    """

    name = "llm"

    @staticmethod
    def is_configured() -> bool:
        # TODO(M-real): 接入 LLM 归一/多模态 OCR 后返回真实判断。
        return False

    async def extract(
        self, *, data: bytes, filename: str, content_type: str, url: str | None
    ) -> ExtractOutput:
        local = await LocalDocExtractor().extract(
            data=data, filename=filename, content_type=content_type, url=url
        )
        # TODO(M-real): local.text 送 LLM 归一为规范 Markdown 后返回。
        return local


def get_extractor() -> DocumentExtractor:
    if get_settings().file_extract_provider == "llm" and LLMExtractor.is_configured():
        return LLMExtractor()
    return LocalDocExtractor()
