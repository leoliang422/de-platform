"""M6.2 图片上传：类型校验、成功落盘、鉴权、工厂回退。"""

from pathlib import Path

import pytest
from httpx import AsyncClient

from app.core.config import get_settings
from app.modules.files.storage import S3Storage, get_storage

# 1x1 透明 PNG
_PNG = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
    b"\x08\x06\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\nIDATx\x9cc\x00\x01"
    b"\x00\x00\x05\x00\x01\r\n-\xb4\x00\x00\x00\x00IEND\xaeB`\x82"
)


async def _token(client: AsyncClient, email: str) -> str:
    await client.post(
        "/auth/register",
        json={"email": email, "password": "secret123", "nickname": email.split("@")[0]},
    )
    resp = await client.post("/auth/login", json={"email": email, "password": "secret123"})
    return resp.json()["access_token"]


def _auth(token: str) -> dict[str, str]:
    return {"Authorization": f"Bearer {token}"}


async def test_upload_image_success(
    client: AsyncClient, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setenv("STORAGE_LOCAL_DIR", str(tmp_path))
    get_settings.cache_clear()
    try:
        token = await _token(client, "up1@test.io")
        resp = await client.post(
            "/files/images",
            headers=_auth(token),
            files={"file": ("a.png", _PNG, "image/png")},
        )
        assert resp.status_code == 201, resp.text
        url = resp.json()["url"]
        assert "/uploads/" in url
        # 文件确实落盘到临时目录
        key = url.split("/uploads/", 1)[1]
        assert (tmp_path / key).exists()
    finally:
        get_settings.cache_clear()


async def test_upload_rejects_non_image(client: AsyncClient) -> None:
    token = await _token(client, "up2@test.io")
    resp = await client.post(
        "/files/images",
        headers=_auth(token),
        files={"file": ("a.txt", b"hello", "text/plain")},
    )
    assert resp.status_code == 400


async def test_upload_requires_auth(client: AsyncClient) -> None:
    resp = await client.post(
        "/files/images",
        files={"file": ("a.png", _PNG, "image/png")},
    )
    assert resp.status_code in (401, 403)


def test_storage_factory_falls_back_to_db(monkeypatch: pytest.MonkeyPatch) -> None:
    assert S3Storage.is_configured() is False
    monkeypatch.setenv("STORAGE_PROVIDER", "s3")
    get_settings.cache_clear()
    try:
        # 凭证未配齐 → 回退到数据库持久化存储（而非临时磁盘 local）
        assert get_storage().name == "db"
    finally:
        get_settings.cache_clear()


async def test_db_storage_persists_and_serves(
    client: AsyncClient, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setenv("STORAGE_PROVIDER", "db")
    get_settings.cache_clear()
    try:
        token = await _token(client, "dbimg@test.io")
        resp = await client.post(
            "/files/images",
            headers=_auth(token),
            files={"file": ("a.png", _PNG, "image/png")},
        )
        assert resp.status_code == 201, resp.text
        url = resp.json()["url"]
        assert "/uploads/" in url
        # 文件已持久化进数据库（stored_files）
        from sqlalchemy import select

        from app.core.database import SessionLocal
        from app.modules.files.models import StoredFile

        key = url.split("/uploads/", 1)[1]
        async with SessionLocal() as db:
            row = (
                await db.execute(select(StoredFile).where(StoredFile.key == key))
            ).scalar_one()
            assert row.content_type == "image/png"
            assert row.data == _PNG
    finally:
        get_settings.cache_clear()


async def test_extract_text_file(client: AsyncClient) -> None:
    token = await _token(client, "ex1@test.io")
    resp = await client.post(
        "/files/extract",
        headers=_auth(token),
        files={"file": ("note.md", b"# hi\ncontent", "text/markdown")},
    )
    assert resp.status_code == 201, resp.text
    body = resp.json()
    assert body["kind"] == "text"
    assert body["placeholder"] is False
    assert "content" in body["text"]


async def test_extract_image_returns_markdown(
    client: AsyncClient, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setenv("STORAGE_LOCAL_DIR", str(tmp_path))
    get_settings.cache_clear()
    try:
        token = await _token(client, "ex2@test.io")
        resp = await client.post(
            "/files/extract",
            headers=_auth(token),
            files={"file": ("p.png", _PNG, "image/png")},
        )
        assert resp.status_code == 201, resp.text
        body = resp.json()
        assert body["kind"] == "image"
        assert body["text"].startswith("![") and "/uploads/" in body["text"]
    finally:
        get_settings.cache_clear()


async def test_extract_docx_real_text(
    client: AsyncClient, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setenv("STORAGE_LOCAL_DIR", str(tmp_path))
    get_settings.cache_clear()
    try:
        # 用 python-docx 生成一个真实 .docx
        from io import BytesIO

        from docx import Document

        doc = Document()
        doc.add_heading("数据倾斜", level=2)
        doc.add_paragraph("倾斜的本质是分区数据分布不均。")
        buf = BytesIO()
        doc.save(buf)

        token = await _token(client, "ex3@test.io")
        docx_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        resp = await client.post(
            "/files/extract",
            headers=_auth(token),
            files={"file": ("r.docx", buf.getvalue(), docx_type)},
        )
        assert resp.status_code == 201, resp.text
        body = resp.json()
        assert body["kind"] == "document"
        assert body["placeholder"] is False  # 真实抽取到文字
        assert "数据倾斜" in body["text"]
        assert "分区数据分布不均" in body["text"]
    finally:
        get_settings.cache_clear()


async def test_extract_invalid_docx_falls_back_to_placeholder(
    client: AsyncClient, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setenv("STORAGE_LOCAL_DIR", str(tmp_path))
    get_settings.cache_clear()
    try:
        token = await _token(client, "ex3b@test.io")
        docx_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        resp = await client.post(
            "/files/extract",
            headers=_auth(token),
            files={"file": ("bad.docx", b"not a real docx", docx_type)},
        )
        assert resp.status_code == 201, resp.text
        body = resp.json()
        assert body["placeholder"] is True  # 解析失败 → 占位 + 下载链接
        assert body["url"] and "/uploads/" in body["url"]
    finally:
        get_settings.cache_clear()


async def test_extract_rejects_unknown_type(client: AsyncClient) -> None:
    token = await _token(client, "ex4@test.io")
    resp = await client.post(
        "/files/extract",
        headers=_auth(token),
        files={"file": ("a.bin", b"\x00\x01", "application/octet-stream")},
    )
    assert resp.status_code == 400
